"""
文件处理服务
负责各种文件格式的文本提取
"""
import os
import tempfile
from pathlib import Path
from typing import Dict, Any, Optional
import PyPDF2
import pypdf
from docx import Document
import aiofiles
from fastapi import HTTPException, UploadFile

from utils.logger import app_logger
from config import settings


class FileProcessor:
    """文件处理器"""
    
    def __init__(self):
        self.supported_types = settings.allowed_file_types
        self.max_size = settings.max_file_size
    
    async def validate_file(self, file: UploadFile) -> None:
        """验证文件"""
        # 检查文件名
        if not file.filename:
            raise HTTPException(status_code=400, detail="文件名不能为空")
        
        # 检查文件类型
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in self.supported_types:
            raise HTTPException(
                status_code=400, 
                detail=f"不支持的文件类型: {file_ext}。支持的类型: {', '.join(self.supported_types)}"
            )
        
        # 检查文件大小
        if file.size and file.size > self.max_size:
            raise HTTPException(
                status_code=400, 
                detail=f"文件大小超过限制: {file.size / 1024 / 1024:.1f}MB > {self.max_size / 1024 / 1024}MB"
            )
        
        app_logger.info(f"文件验证通过: {file.filename}, 大小: {file.size} bytes")
    
    async def extract_text_from_pdf(self, file_path: str) -> str:
        """从PDF文件提取文本"""
        try:
            app_logger.info(f"开始提取PDF文本: {file_path}")
            text = ""
            
            with open(file_path, 'rb') as file:
                # 优先使用pypdf（更现代）
                try:
                    pdf_reader = pypdf.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"=== 第{page_num + 1}页 ===\n{page_text}\n\n"
                except Exception as e:
                    app_logger.warning(f"pypdf提取失败，尝试PyPDF2: {e}")
                    # 回退到PyPDF2
                    file.seek(0)
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text += f"=== 第{page_num + 1}页 ===\n{page_text}\n\n"
            
            extracted_text = text.strip()
            app_logger.info(f"PDF文本提取完成，提取了 {len(extracted_text)} 个字符")
            return extracted_text
            
        except Exception as e:
            app_logger.error(f"PDF文本提取失败: {e}")
            raise HTTPException(status_code=400, detail=f"PDF文件解析失败: {str(e)}")
    
    async def extract_text_from_docx(self, file_path: str) -> str:
        """从DOCX文件提取文本"""
        try:
            app_logger.info(f"开始提取DOCX文本: {file_path}")
            doc = Document(file_path)
            text = ""
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            extracted_text = text.strip()
            app_logger.info(f"DOCX文本提取完成，提取了 {len(extracted_text)} 个字符")
            return extracted_text
            
        except Exception as e:
            app_logger.error(f"DOCX文本提取失败: {e}")
            raise HTTPException(status_code=400, detail=f"DOCX文件解析失败: {str(e)}")
    
    async def extract_text_from_txt(self, file_path: str) -> str:
        """从TXT文件提取文本"""
        try:
            app_logger.info(f"开始提取TXT文本: {file_path}")
            
            # 尝试不同编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            content = None
            
            for encoding in encodings:
                try:
                    async with aiofiles.open(file_path, 'r', encoding=encoding) as file:
                        content = await file.read()
                    app_logger.info(f"使用编码 {encoding} 成功读取文件")
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError("无法使用任何编码读取文件")
            
            extracted_text = content.strip()
            app_logger.info(f"TXT文本提取完成，提取了 {len(extracted_text)} 个字符")
            return extracted_text
            
        except Exception as e:
            app_logger.error(f"TXT文本提取失败: {e}")
            raise HTTPException(status_code=400, detail=f"TXT文件解析失败: {str(e)}")
    
    async def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """根据文件类型提取文本"""
        file_type = file_type.lower()
        
        if file_type == '.pdf':
            return await self.extract_text_from_pdf(file_path)
        elif file_type in ['.docx', '.doc']:
            return await self.extract_text_from_docx(file_path)
        elif file_type in ['.txt', '.md']:
            return await self.extract_text_from_txt(file_path)
        else:
            raise HTTPException(status_code=400, detail=f"不支持的文件类型: {file_type}")
    
    async def process_uploaded_file(self, file: UploadFile) -> Dict[str, Any]:
        """处理上传的文件"""
        try:
            # 验证文件
            await self.validate_file(file)
            
            # 获取文件信息
            file_ext = Path(file.filename).suffix.lower()
            file_size = file.size or 0
            
            # 创建临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
                content = await file.read()
                temp_file.write(content)
                temp_file_path = temp_file.name
            
            try:
                # 提取文本内容
                extracted_content = await self.extract_text_from_file(temp_file_path, file_ext)
                
                # 限制内容长度
                if len(extracted_content) > settings.max_content_length:
                    extracted_content = extracted_content[:settings.max_content_length]
                    app_logger.warning(f"文件内容过长，已截取前 {settings.max_content_length} 个字符")
                
                return {
                    "content": extracted_content,
                    "filename": file.filename,
                    "size": file_size,
                    "type": file_ext,
                    "extractedLength": len(extracted_content)
                }
                
            finally:
                # 清理临时文件
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    app_logger.debug(f"已清理临时文件: {temp_file_path}")
        
        except HTTPException:
            raise
        except Exception as e:
            app_logger.error(f"文件处理失败: {e}")
            raise HTTPException(status_code=500, detail=f"文件处理失败: {str(e)}")


# 全局文件处理器实例
file_processor = FileProcessor()

